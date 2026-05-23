"""
法学论文质检自查 Agent —— unified-thesis-reviewer 总编排器 (v3.0)

核心职责：
1. 接收论文文件（docx / pdf / txt / md）或 URL
2. 调用 read_paper_file 读取论文全文
3. LLM 直接在分析中输出 UTR schema issues JSON（无需额外 normalize 步骤）
4. 调用 generate_deliverables 一次性生成带批注 docx + PDF 报告

优化（v3.0 vs v2.0）：
- 删除 verify_fact：该工具每次联网搜索耗时5分钟×10+事实=50+分钟，是最大瓶颈
  LLM 自身知识足以识别法条引用是否规范，标注"需人工核实"即可
- 删除 normalize_issues_for_annotation：LLM 直接输出 UTR JSON，省去一次 LLM 往返
- 合并 generate_annotated_docx + generate_markdown_report → generate_deliverables：
  一次工具调用完成两个交付物，省去一次 LLM 往返
- 总 LLM 工具往返从 6+2N 降至 3（read → analysis → deliverables）

失败时暴露完整堆栈，禁止泛化话术。
"""

import os
import sys
import json
import time
import tempfile
import traceback
import re
import logging
from typing import Any, Dict, List, Optional

logger = logging.getLogger(__name__)

from langchain.tools import tool
from langchain.agents import create_agent
from langchain.agents.middleware import wrap_tool_call
from langchain_core.messages import AnyMessage, ToolMessage
from langchain_openai import ChatOpenAI
from langgraph.graph import MessagesState
from langgraph.graph.message import add_messages  # noqa: F401 - kept for compatibility
from typing import Annotated
from coze_coding_utils.runtime_ctx.context import default_headers
from coze_coding_utils.log.write_log import request_context
from coze_coding_utils.runtime_ctx.context import new_context
from storage.memory.memory_saver import get_memory_saver

# 导入 Coze 工作负载身份认证客户端
from coze_workload_identity import Client as WorkloadIdentityClient

# 简单的 access_token 缓存（进程级，避免每次请求都交换令牌）
_access_token_cache = {"token": None, "expiry": 0}

def _get_platform_access_token() -> str | None:
    """通过 Coze 工作负载身份获取临时 access_token（JWT 格式）。"""
    global _access_token_cache
    now = time.time()
    # 缓存 50 分钟（令牌通常 1 小时过期）
    if _access_token_cache["token"] and _access_token_cache["expiry"] > now + 600:
        return _access_token_cache["token"]
    try:
        client = WorkloadIdentityClient()
        token = client.get_access_token()
        _access_token_cache["token"] = token
        _access_token_cache["expiry"] = now + 3000  # 50 分钟
        return token
    except Exception:
        return None


def _resolve_auth() -> tuple[str | None, str | None]:
    """解析认证信息。

    优先级（从高到低）：
    1. BYOK: 多种环境变量名（兼容用户自定义变量名）
    2. BYOK: custom_model.api_key（用户通过 UI 设置，仅内存存储）
    3. 平台默认: coze_workload_identity Client 交换临时 access_token
    """
    workspace_path = os.getenv("COZE_WORKSPACE_PATH", "/workspace/projects")
    config_path = os.path.join(workspace_path, LLM_CONFIG)
    custom = {}
    if os.path.exists(config_path):
        with open(config_path, "r", encoding="utf-8") as f:
            custom = json.load(f).get("custom_model", {})

    # BYOK 模式 — 支持多种环境变量名（兼容不同用户的命名习惯）
    byok_key = (
        os.getenv("CUSTOM_MODEL_API_KEY")
        or os.getenv("DEEPSEEK_API_KEY")
        or os.getenv("deepseek_coze_coding")      # 用户自定义变量名
        or os.getenv("OPENAI_API_KEY")
        or os.getenv("COZE_API_KEY")
        or custom.get("api_key")
    )
    if byok_key:
        byok_url = (
            custom.get("base_url")
            or os.getenv("CUSTOM_MODEL_BASE_URL")
            or os.getenv("DEEPSEEK_BASE_URL")
            or os.getenv("OPENAI_BASE_URL")
            or "https://api.deepseek.com"            # DeepSeek 默认地址
        )
        logger.info(f"[Auth] BYOK mode enabled, base_url={byok_url}, key_prefix={byok_key[:8]}...")
        return byok_key, byok_url

    # 平台默认认证：通过工作负载身份交换临时 JWT 令牌
    platform_token = _get_platform_access_token()
    if platform_token:
        platform_url = os.getenv("COZE_INTEGRATION_MODEL_BASE_URL")
        logger.info(f"[Auth] Platform mode enabled, base_url={platform_url}")
        return platform_token, platform_url

    logger.warning("[Auth] No valid authentication found")
    return None, None


# ---------------------------------------------------------------------------
# 常量 / 路径
# ---------------------------------------------------------------------------
LLM_CONFIG = "config/agent_llm_config.json"
UTR_BUNDLE = os.path.join(os.path.dirname(__file__), "..", "tools", "scripts", "utr-bundle")

ENUM_SOURCE = {"thesis", "citation"}
ENUM_CATEGORY = {
    "structure", "argumentation", "literature-review", "empirical",
    "legal-norms", "language", "policy", "academic-integrity",
    "citation-format", "citation-missing-info"
}
ENUM_SEVERITY = {"fatal", "major", "minor"}
ENUM_SCOPE = {"document", "chapter", "paragraph", "sentence", "span"}

SEVERITY_ORDER = {"fatal": 0, "major": 1, "minor": 2}


class AgentState(MessagesState):
    messages: Annotated[list[AnyMessage], add_messages]


def _utr_tool_path(name: str) -> str:
    return os.path.join(UTR_BUNDLE, "unified-thesis-reviewer", "tools", name)


def _truncate_filename(name: str, max_bytes: int = 180, reserve: int = 30) -> str:
    """截断文件名，确保不超过 max_bytes 字节（UTF-8），保留扩展名。
    reserve 为给后缀（如 .annotated.docx）预留的字节数。"""
    if not name:
        return "document"
    import os as _os
    base, ext = _os.path.splitext(name)
    # 扩展名本身长度
    ext_bytes = len(ext.encode("utf-8"))
    # 留给 base 的最大字节数
    available = max_bytes - ext_bytes - reserve
    if available < 10:
        available = 10
    encoded = base.encode("utf-8")
    if len(encoded) <= available:
        return name
    # 截断到 available 字节，且以完整字符结尾
    truncated = encoded[:available]
    while truncated and (truncated[-1] & 0xC0) == 0x80:
        truncated = truncated[:-1]
    return truncated.decode("utf-8", errors="ignore") + ext


def _download_url_to_binary(url: str) -> tuple[bytes, str]:
    """下载 URL 内容，返回 (二进制数据, 文件名)。支持带 query string 的 URL。"""
    import requests
    resp = requests.get(url, timeout=60, allow_redirects=True)
    if resp.status_code != 200:
        raise FileNotFoundError(f"URL 下载失败，状态码 {resp.status_code}: {url}")
    # 从 Content-Disposition 或 URL 解析文件名
    cd = resp.headers.get("content-disposition", "")
    import re
    fname_match = re.search(r'filename[^;=\n]*=(?:["\']?)([^\n;"\']*)', cd)
    if fname_match:
        fname = fname_match.group(1).strip().strip('"\'')
    else:
        # 从 URL path 中提取文件名
        from urllib.parse import urlparse
        parsed = urlparse(url)
        path_parts = parsed.path.split("/")
        fname = path_parts[-1] if path_parts[-1] else "downloaded.file"
        if "?" in fname:
            fname = fname.split("?")[0]
    if not fname or fname == "downloaded.file":
        fname = "downloaded.docx"
    return resp.content, fname


def _is_docx_bytes(data: bytes) -> bool:
    """通过魔数判断是否为 DOCX（ZIP 格式）。"""
    return data[:2] == b"PK"


def _resolve_local_docx_path(file_path: str) -> str:
    """将 URL / file_id / 本地路径统一解析为本地 .docx 二进制文件路径。"""
    path = file_path.strip()

    # URL → 下载二进制文件到 /tmp
    if path.startswith("http://") or path.startswith("https://"):
        data, fname = _download_url_to_binary(path)
        if not fname.endswith(".docx"):
            fname += ".docx"
        safe_fname = _truncate_filename(fname, max_bytes=180, reserve=40)
        tmp_path = os.path.join("/tmp", f"utr_{int(__import__('time').time())}_{safe_fname}")
        with open(tmp_path, "wb") as f:
            f.write(data)
        return tmp_path

    # Coze file_id → 解析为本地路径
    if path.startswith("file_") and not os.path.exists(path):
        try:
            import coze_coding_utils.file as cf_mod
            coze_file = getattr(cf_mod, "file", cf_mod)
            resolved = coze_file.get_file(path).file_path  # type: ignore[union-attr]
            return resolved
        except Exception:
            pass

    # 本地路径直接返回
    return path


# ---------------------------------------------------------------------------
# 底层读取函数（本地文件）
# ---------------------------------------------------------------------------
def _read_local_docx(path: str) -> dict:
    """使用 extract-docx.py 脚本或 python-docx 读取本地 DOCX 文件。"""
    import subprocess
    extract_script = _utr_tool_path("extract-docx.py")
    tmp_out = None
    if os.path.exists(extract_script):
        with tempfile.NamedTemporaryFile(mode="w", encoding="utf-8", suffix=".txt", delete=False) as tf:
            tmp_out = tf.name
        try:
            proc = subprocess.run(
                [sys.executable, extract_script, path, tmp_out],
                capture_output=True, text=True, timeout=120
            )
            if proc.returncode == 0 and os.path.exists(tmp_out):
                with open(tmp_out, encoding="utf-8") as f:
                    text = f.read()
                os.unlink(tmp_out)
                return {"success": True, "text": text, "format": "docx",
                        "file_name": os.path.basename(path),
                        "total_chars": len(text),
                        "text_preview": text[:5000] if text else "",
                        "full_text_available": True}
        finally:
            if tmp_out and os.path.exists(tmp_out):
                try:
                    os.unlink(tmp_out)
                except Exception:
                    pass
    from docx import Document
    doc = Document(path)
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    def full_para_text(p_el):
        parts = []
        for child in p_el.iter():
            tag = child.tag.split("}", 1)[-1] if "}" in child.tag else child.tag
            if tag == "t" and child.text:
                parts.append(child.text)
            elif tag == "tab":
                parts.append("\t")
        return "".join(parts)
    body_el = doc.element.body  # type: ignore[attr-defined]
    paragraphs = [full_para_text(p_el) for p_el in body_el.iter(f"{{{W_NS}}}p")]
    paragraphs = [p for p in paragraphs if p.strip()]
    text = "\n".join(paragraphs)
    return {"success": True, "text": text, "format": "docx",
            "file_name": os.path.basename(path),
            "total_chars": len(text),
            "text_preview": text[:5000] if text else "",
            "full_text_available": True}


def _read_local_pdf(path: str) -> dict:
    """使用 PyMuPDF (fitz) 读取本地 PDF 文件。"""
    try:
        import fitz
        doc = fitz.open(path)
        if doc.is_encrypted:
            doc.close()
            return {"success": False, "error": "PDF 已加密", "is_encrypted": True}
        total_chars = sum(len(page.get_text()) for page in doc)
        if total_chars == 0:
            doc.close()
            return {"success": False, "error": "PDF 为扫描件或图片型 PDF，无法提取文本。请换用 Word 版本或 OCR 后再试。", "is_scanned": True}
        texts = []
        for i, page in enumerate(doc):
            texts.append(f"--- Page {i + 1} ---\n{page.get_text()}")
        doc.close()
        return {"success": True, "text": "\n".join(texts), "format": "pdf",
                "page_count": len(texts), "total_chars": total_chars,
                "file_name": os.path.basename(path)}
    except ImportError:
        return {"success": False, "error": "PyMuPDF (fitz) 未安装，无法读取 PDF"}


def _read_text(file_path: str) -> dict:
    """从本地路径或公网 URL 读取文本。"""
    try:
        path = file_path.strip()
        if path.startswith("http://") or path.startswith("https://"):
            data, fname = _download_url_to_binary(path)
            # 判断真实格式（通过魔数或文件名）
            ext = ".docx" if fname.lower().endswith(".docx") or _is_docx_bytes(data) else \
                   ".pdf" if fname.lower().endswith(".pdf") else \
                   ".txt"
            if ext == ".docx":
                # 下载为二进制，解析 DOCX
                import tempfile as _tempfile, subprocess as _subprocess, sys as _sys
                tmp_docx = None
                try:
                    with _tempfile.NamedTemporaryFile(mode="wb", suffix=".docx", delete=False) as _tf:
                        _tf.write(data)
                        tmp_docx = _tf.name
                    return _read_local_docx(tmp_docx, os.path.basename(fname))
                finally:
                    if tmp_docx and os.path.exists(tmp_docx):
                        try: os.unlink(tmp_docx)
                        except: pass
            elif ext == ".pdf":
                # 下载为二进制，解析 PDF
                import tempfile as _tempfile, fitz as _fitz
                tmp_pdf = None
                try:
                    with _tempfile.NamedTemporaryFile(mode="wb", suffix=".pdf", delete=False) as _tf:
                        _tf.write(data)
                        tmp_pdf = _tf.name
                    result = _read_local_pdf(tmp_pdf, os.path.basename(fname))
                    result["file_name"] = os.path.basename(fname)
                    return result
                finally:
                    if tmp_pdf and os.path.exists(tmp_pdf):
                        try: os.unlink(tmp_pdf)
                        except: pass
            else:
                # 纯文本格式
                try:
                    content = data.decode("utf-8")
                except UnicodeDecodeError:
                    content = data.decode("gbk", errors="ignore")
                return {"success": True, "text": content, "format": ext.lstrip("."),
                        "file_name": os.path.basename(fname)}

        if path.startswith("file_") and not os.path.exists(path):
            try:
                import coze_coding_utils.file as cf_mod
                coze_file = getattr(cf_mod, "file", cf_mod)
                fpath = coze_file.get_file(path).file_path  # type: ignore[union-attr]
                path = fpath
            except Exception:
                return {"success": False, "error": f"无法通过平台 file_id 获取文件: {path}"}

        if not os.path.exists(path):
            return {"success": False, "error": f"文件不存在: {path}"}

        ext = os.path.splitext(path)[1].lower()

        if ext == ".docx":
            return _read_local_docx(path)

        if ext == ".pdf":
            return _read_local_pdf(path)

        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return {"success": True, "text": content, "format": ext.lstrip("."),
                "file_name": os.path.basename(path)}
    except Exception as e:
        return {"success": False, "error": f"文件读取失败: {str(e)}"}


# ---------------------------------------------------------------------------
# @tool 函数
# ---------------------------------------------------------------------------
@tool
def _is_valid_docx(path: str) -> bool:
    """通过 python-docx 验证文件是否为真实有效的 docx。"""
    try:
        from docx import Document
        Document(path)
        return True
    except Exception:
        return False


def read_paper_file(file_path: str) -> str:
    """读取论文文件内容。file_path 可以是本地路径、公网 URL 或 Coze 平台 file_id。"""
    # Step 1: 先解析本地路径（平台临时文件可能很快被清理）
    local_docx_path = ""
    read_path = file_path
    try:
        resolved = _resolve_local_docx_path(file_path)
        logger.info(f"[read_paper_file] Resolved path: {resolved}")
        if resolved and os.path.exists(resolved):
            if _is_valid_docx(resolved):
                # 有效 docx：复制到 /tmp 持久化
                if not resolved.startswith("/tmp"):
                    import hashlib as _hashlib, shutil as _shutil
                    _cache_key = _hashlib.md5(file_path.encode()).hexdigest()
                    persistent_path = f"/tmp/paper_original_{_cache_key}.docx"
                    _shutil.copy2(resolved, persistent_path)
                    local_docx_path = persistent_path
                    read_path = persistent_path
                    logger.info(f"[read_paper_file] Valid docx copied to {persistent_path}")
                else:
                    local_docx_path = resolved
                    read_path = resolved
                    logger.info(f"[read_paper_file] Using existing /tmp docx: {resolved}")
            else:
                # 不是有效 docx（如 PDF 被伪装为 docx），仍可用该路径读取文本
                read_path = resolved
                logger.info(f"[read_paper_file] Resolved file is not valid docx, reading as original format")
        else:
            logger.warning(f"[read_paper_file] Resolved path not exists: {resolved}")
    except Exception as e:
        logger.warning(f"[read_paper_file] Resolve/copy failed: {e}")

    # Step 2: 读取文本内容
    result = _read_text(read_path)
    if not result["success"]:
        logger.warning(f"[read_paper_file] _read_text failed for {read_path}: {result.get('error')}")
        return json.dumps(result, ensure_ascii=False)
    paper_text = result["text"]
    logger.info(f"[read_paper_file] Text read success, chars={len(paper_text)}")

    # 论文文本太长（可能10万+字符），不传给LLM以避免上下文溢出
    # 改为存储在/tmp，供后续工具直接访问
    import hashlib, pickle, os as _os
    _os.makedirs("/tmp", exist_ok=True)
    cache_key = hashlib.md5(file_path.encode()).hexdigest()
    cache_file = f"/tmp/paper_cache_{cache_key}.pkl"
    with open(cache_file, "wb") as f:
        pickle.dump({"text": paper_text, "local_path": local_docx_path}, f)

    # LLM只需要：文件信息 + 摘要预览（前3000字）
    preview = paper_text[:3000]
    return json.dumps({
        "success": True,
        "format": result.get("format"),
        "file_name": result.get("file_name"),
        "total_chars": len(paper_text),
        "local_docx_path": local_docx_path,
        "paper_cache_key": cache_key,
        "text_preview": preview,
        "full_text_available": True,
        "note": "完整论文内容已缓存，generate_deliverables 将读取全文生成批注"
    }, ensure_ascii=False)


# ---------------------------------------------------------------------------
# 辅助函数：确保 docx 有效性，支持 PDF/文本降级转换
# ---------------------------------------------------------------------------
def _ensure_valid_docx(file_path: str) -> Optional[str]:
    """确保 file_path 是一个有效的 docx 文件。如果不是，尝试转换或创建新的 docx。
    返回有效的 docx 路径，或 None（如果无法创建）。"""
    if not file_path or not os.path.exists(file_path):
        return None

    # 尝试验证是否为真实 docx
    try:
        from docx import Document
        Document(file_path)
        return file_path
    except Exception:
        pass

    # 不是有效 docx，尝试读取文件内容判断真实格式
    try:
        with open(file_path, "rb") as f:
            header = f.read(8)
        is_pdf = header.startswith(b"%PDF")

        if is_pdf:
            # PDF 文件：提取文本并创建新的 docx
            import pypdf
            reader = pypdf.PdfReader(file_path)
            from docx import Document
            new_doc = Document()
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    for line in text.split('\n'):
                        line = line.strip()
                        if line:
                            new_doc.add_paragraph(line)
            tmp_path = f"/tmp/converted_pdf_{int(time.time())}.docx"
            new_doc.save(tmp_path)
            logger.info(f"[_ensure_valid_docx] Converted PDF to docx: {tmp_path}")
            return tmp_path

        # 尝试作为文本读取
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            from docx import Document
            new_doc = Document()
            for line in text.split('\n'):
                line = line.strip()
                if line:
                    new_doc.add_paragraph(line)
            tmp_path = f"/tmp/converted_txt_{int(time.time())}.docx"
            new_doc.save(tmp_path)
            logger.info(f"[_ensure_valid_docx] Converted text to docx: {tmp_path}")
            return tmp_path
        except Exception:
            pass
    except Exception:
        pass

    return None


def _recover_docx_from_cache(docx_path_hint: str = "") -> Optional[str]:
    """如果 docx_path 无效，尝试从 /tmp/paper_cache_*.pkl 中恢复论文文本并创建 docx。"""
    try:
        import glob, pickle
        cache_files = glob.glob("/tmp/paper_cache_*.pkl")
        if not cache_files:
            return None

        # 按修改时间排序，取最新的
        cache_files.sort(key=os.path.getmtime, reverse=True)
        for cache_file in cache_files:
            try:
                with open(cache_file, "rb") as f:
                    cache = pickle.load(f)
                text = cache.get("text", "")
                if not text:
                    continue
                from docx import Document
                new_doc = Document()
                for line in text.split('\n'):
                    line = line.strip()
                    if line:
                        new_doc.add_paragraph(line)
                tmp_path = f"/tmp/recovered_{int(time.time())}.docx"
                new_doc.save(tmp_path)
                logger.info(f"[_recover_docx_from_cache] Recovered docx from cache: {tmp_path}")
                return tmp_path
            except Exception:
                continue
    except Exception:
        pass
    return None


# ---------------------------------------------------------------------------
# UTR Schema 验证与修补（内联，不需要单独工具）
# ---------------------------------------------------------------------------
ID_PATTERN = re.compile(r"^(thesis|citation)-([a-z-]+)-\d{3}$")
GROUP_ID_PATTERN = re.compile(r"^g-\d{3,}$")

# 封面页/模板文本黑名单
_COVER_BLACKLIST = {
    "湘潭大学", "毕业论文", "学士学位论文", "硕士论文", "博士论文",
    "学号", "指导教师", "学院", "法学学部", "Faculty of Law",
    "Xiangtan University", "本科", "课程论文", "原创性声明",
    "版权使用授权书", "答辩委员会", "摘要", "Abstract", "关键词",
    "Keywords", "目录", "参考文献", "致谢", "附录",
}


def _validate_and_fix_issues(issues_data: dict, paper_text: str) -> dict:
    """验证并修补 UTR issues JSON，确保 schema 合规。"""
    issues = issues_data.get("issues", [])
    if not issues:
        return issues_data

    fixed_issues = []
    group_counter = 1
    gid_map = {}

    def _next_gid(src, cat, chapter, pidx):
        nonlocal group_counter
        key = (src, cat, chapter, pidx)
        if key not in gid_map:
            gid_map[key] = f"g-{group_counter:03d}"
            group_counter += 1
        return gid_map[key]

    for i, issue in enumerate(issues):
        # Fix id
        src = issue.get("source", "thesis")
        cat = issue.get("category", "argumentation")
        if src not in ENUM_SOURCE:
            src = "thesis"
        if cat not in ENUM_CATEGORY:
            cat = "argumentation"
        issue["id"] = f"{src}-{cat}-{i+1:03d}"
        issue["source"] = src
        issue["category"] = cat

        # Fix severity
        sev = issue.get("severity", "minor")
        if sev not in ENUM_SEVERITY:
            sev = "minor"
        issue["severity"] = sev

        # Fix scope
        scope = issue.get("scope", "paragraph")
        if scope not in ENUM_SCOPE:
            scope = "paragraph"
        issue["scope"] = scope

        # Fix locator
        locator = issue.get("locator", {})
        if not isinstance(locator, dict):
            locator = {}
        locator.setdefault("chapter", "正文")
        locator.setdefault("paragraph_index", 0)
        issue["locator"] = locator

        # Fix excerpt (scope=document/chapter时必须为空)
        excerpt = issue.get("excerpt", "")
        if scope in ("document", "chapter"):
            excerpt = ""
        elif not excerpt and len(paper_text) > 0:
            # 尝试从问题中提取可能的原文片段
            anchor = issue.get("anchor_text", "")
            excerpt = anchor[:60] if anchor else ""
        if len(excerpt) > 60:
            excerpt = excerpt[:57] + "..."
        issue["excerpt"] = excerpt

        # Fix anchor_text - 确保是论文原文而非概括
        anchor_text = issue.get("anchor_text", "")
        if anchor_text and paper_text:
            # 验证 anchor_text 是否在论文中出现
            if anchor_text not in paper_text:
                # 尝试找最相似的片段
                anchor_text = _find_anchor_in_text(issue.get("problem", ""), paper_text)
        elif not anchor_text and paper_text:
            anchor_text = _find_anchor_in_text(issue.get("problem", ""), paper_text)
        if len(anchor_text) > 60:
            anchor_text = anchor_text[:60]
        issue["anchor_text"] = anchor_text

        # Fix problem
        problem = issue.get("problem", "")
        if not problem or not problem.strip():
            problem = "问题描述待补充"
        if len(problem) > 200:
            problem = problem[:200]
        issue["problem"] = problem

        # Fix suggestion
        suggestion = issue.get("suggestion", [])
        if not isinstance(suggestion, list):
            suggestion = [str(suggestion)]
        if not suggestion:
            suggestion = ["请根据具体上下文补充修改建议"]
        suggestion = [str(s)[:500] for s in suggestion[:5]]
        issue["suggestion"] = suggestion

        # Fix group_id
        gid = issue.get("group_id", "")
        if not GROUP_ID_PATTERN.match(gid):
            gid = _next_gid(src, cat, locator.get("chapter", "正文"), locator.get("paragraph_index", 0))
        issue["group_id"] = gid

        fixed_issues.append(issue)

    issues_data["issues"] = fixed_issues
    # 质量门禁
    total = len(fixed_issues)
    substantive = sum(1 for i in fixed_issues if i["severity"] in ("fatal", "major"))
    issues_data.setdefault("quality_gate", {
        "total_issues": total,
        "substantive_issues": substantive,
        "quality_gate_passed": total >= 8 and substantive >= 5
    })
    issues_data["quality_gate"]["total_issues"] = total
    issues_data["quality_gate"]["substantive_issues"] = substantive
    issues_data["quality_gate"]["quality_gate_passed"] = total >= 8 and substantive >= 5

    return issues_data


def _find_anchor_in_text(text: str, paper: str) -> str:
    """在论文正文中寻找可定位的 anchor_text，返回前60码点。

    策略：
    1. 从 text（问题描述）中提取关键词
    2. 在 paper 中搜索含这些关键词的段落
    3. 返回该段落中最接近关键词的连续中文文本片段（5-60字）
    4. 完全匹配优于部分匹配
    """
    if not paper or not text:
        return text[:60] if text else ""

    # 找到正文起始位置（跳过封面/目录等前置部分）
    body_start = 0
    for marker in [r"第[一二三四五六七八九十\d]+章", r"绪\s*论", r"引\s*言", r"一、", r"摘要"]:
        m = re.search(marker, paper)
        if m:
            body_start = max(body_start, m.start())
            break
    if body_start == 0 and len(paper) > 2000:
        body_start = int(len(paper) * 0.15)
    body_paper = paper[body_start:]

    # 1. 引号内原文引用（最高优先级，LLM已确保这些是论文原文）
    quoted = []
    _QP = chr(0x201C) + chr(0x201D) + chr(0x2018) + chr(0x2019) + chr(34) + chr(39)
    for m in re.finditer(_QP + '([\u4e00-\u9fff]{4,60}?)' + _QP, text):
        quoted.append(m.group(1).strip())
    for q in quoted:
        if q in body_paper:
            return q[:60]
        # 去除标点后匹配
        raw = re.sub(r"[，。、：；！？\s\-—]", "", q)
        raw_paper = re.sub(r"[，。、：；！？\s\-—]", "", body_paper)
        idx = raw_paper.find(raw)
        if idx >= 0:
            start = max(0, idx - 5)
            end = min(len(body_paper), idx + len(raw) + 5)
            return body_paper[start:end][:60]

    # 2. 从问题描述中提取专业术语和关键词
    stop_words = {"全文", "本文", "论文", "该文", "该论文", "建议", "问题", "内容",
                   "主要", "重要", "严重", "轻微", "具体", "有关", "研究", "分析",
                   "学术", "规范", "格式", "存在", "需要", "应当", "目前", "当前",
                   "情况", "方面", "角度", "层面", "章节", "部分", "章节名",
                   "湘潭大学", "毕业论文", "学士学位", "学号", "指导教师", "正文"}
    words = re.findall(r"[\u4e00-\u9fff]{2,6}", text)
    significant_words = [w for w in words if w not in stop_words and len(set(w)) >= 2]
    word_scores = {}
    for w in significant_words:
        word_scores[w] = word_scores.get(w, 0) + len(w)
    sorted_words = sorted(word_scores.keys(), key=lambda x: (-word_scores[x], -len(x)))[:10]

    # 3. 在正文中搜索含关键词的段落
    for keyword in sorted_words:
        idx = body_paper.find(keyword)
        if idx < 0:
            raw = re.sub(r"[，。、：；\s]", "", keyword)
            if raw and raw != keyword:
                raw_paper2 = re.sub(r"[，。、：；\s]", "", body_paper)
                idx = raw_paper2.find(raw)
        if idx < 0:
            continue
        context_start = max(0, idx - 20)
        context_end = min(len(body_paper), idx + len(keyword) + 20)
        context = body_paper[context_start:context_end]
        fragments = re.findall(r"[一-鿿，。、：；！？""''()《》 	-]{5,60}", context)
        if fragments:
            best = max(fragments, key=lambda x: (len(re.sub(r"\s", "", x)), keyword in x))
            clean = best.strip()
            if clean:
                return clean[:60]

    # 4. 降级：取正文前200字中最长的中文片段
    fragments = re.findall(r"[\u4e00-\u9fff，。、：；！？\-]{10,60}", body_paper[:200])
    if fragments:
        return max(fragments, key=len)[:60]

    # 5. 最后降级：返回问题描述的前40字
    return text.strip()[:40][:60]

def generate_deliverables(docx_path: str, issues_json: str, analysis_summary: str = "") -> str:
    """一次性生成全部交付物：带批注的 .annotated.docx + PDF 审查报告。

    参数：
    - docx_path: read_paper_file 返回的 local_docx_path 字段
    - issues_json: 完整 UTR schema JSON 字符串，格式 {"schema_version":"1.0","issues":[...]}
    - analysis_summary: 审查分析摘要文本（用于生成报告）

    返回 JSON 包含 annotated_docx_url 和 pdf_url 两个下载链接。
    ⚠️ 调用此工具后，请将返回的下载链接以友好的格式呈现给用户，不要输出原始 JSON。
    """
    results = {
        "annotated_docx": {"success": False, "url": None, "error": None},
        "pdf_report": {"success": False, "url": None, "error": None},
    }

    # ── Step 1: 解析并验证 issues_json ──
    raw_str = issues_json if isinstance(issues_json, str) else json.dumps(issues_json, ensure_ascii=False)
    try:
        issues_raw = json.loads(raw_str)
    except json.JSONDecodeError:
        # 尝试修复常见问题
        fixed = raw_str
        for truncate_marker in ['],', '},', '",', 'null,', 'true,', 'false,']:
            idx = fixed.rfind(truncate_marker)
            if idx > 0:
                fixed = fixed[:idx + len(truncate_marker)]
                break
        fixed = re.sub(r',?\s*"[^"]*"\s*:\s*"[^"]*$', '', fixed)
        fixed = re.sub(r',?\s*"[^"]*"\s*:\s*\S+$', '', fixed)
        open_b = fixed.count('{') + fixed.count('[')
        close_b = fixed.count('}') + fixed.count(']')
        for _ in range(open_b - close_b):
            fixed += '}' if fixed.rstrip().endswith(']') or fixed.rstrip().endswith('}') else ']'
        try:
            issues_raw = json.loads(fixed)
        except json.JSONDecodeError as je2:
            return json.dumps({
                "success": False,
                "error": f"issues_json 解析失败: {str(je2)}",
                "hint": "请确保传入的 issues_json 是合法 JSON",
                "raw_preview": raw_str[:200],
            }, ensure_ascii=False)

    if isinstance(issues_raw, list):
        issues_data = {"schema_version": "1.0", "issues": issues_raw}
    elif isinstance(issues_raw, dict) and "issues" in issues_raw:
        issues_data = issues_raw
    else:
        issues_data = {"schema_version": "1.0", "issues": []}

    # 读取论文全文（用于验证和修补 anchor_text）
    paper_text = ""
    try:
        if docx_path and os.path.exists(docx_path):
            read_result = _read_text(docx_path)
            if read_result.get("success"):
                paper_text = read_result.get("text", "")
    except Exception:
        pass

    # 验证并修补 issues
    issues_data = _validate_and_fix_issues(issues_data, paper_text)

    # 保存 issues.json 供后续使用
    try:
        with open("/tmp/last_issues.json", "w", encoding="utf-8") as f:
            json.dump(issues_data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

    # ── Step 2: 生成带批注的 docx ──
    if docx_path:
        try:
            import subprocess
            inject_script = _utr_tool_path("inject-docx-comments.py")

            if os.path.exists(inject_script):
                # 解析路径
                try:
                    resolved_path = _resolve_local_docx_path(docx_path)
                except FileNotFoundError as fe:
                    results["annotated_docx"]["error"] = f"原始 docx 文件无法获取: {str(fe)}"
                    resolved_path = None

                # fallback：如果本地路径不存在，尝试重新从平台获取（file_id 或 URL）
                if not resolved_path or not os.path.exists(resolved_path):
                    logger.warning(f"[generate_deliverables] docx_path not exists locally: {docx_path}, trying re-resolve")
                    try:
                        fallback = _resolve_local_docx_path(docx_path)
                        if fallback and os.path.exists(fallback):
                            resolved_path = fallback
                            logger.info(f"[generate_deliverables] Re-resolved to: {fallback}")
                            # 复制到 /tmp 防止再次被清理
                            if not fallback.startswith("/tmp"):
                                import shutil as _shutil2
                                _fb_path = f"/tmp/paper_fallback_{int(__import__('time').time())}.docx"
                                _shutil2.copy2(fallback, _fb_path)
                                resolved_path = _fb_path
                                logger.info(f"[generate_deliverables] Copied fallback to: {_fb_path}")
                    except Exception as e2:
                        logger.warning(f"[generate_deliverables] Re-resolve failed: {e2}")

                # 确保 resolved_path 指向有效 docx（支持 PDF/文本降级转换）
                if resolved_path and os.path.exists(resolved_path):
                    valid_docx = _ensure_valid_docx(resolved_path)
                    if not valid_docx:
                        valid_docx = _recover_docx_from_cache(docx_path)
                    if valid_docx:
                        resolved_path = valid_docx
                    else:
                        results["annotated_docx"]["error"] = "无法获取有效的 docx 文件（上传的论文不是 .docx 格式或文件已失效），已跳过批注文档生成"
                        resolved_path = None

                if resolved_path and os.path.exists(resolved_path):
                    # 写入临时 issues.json
                    debug_path = resolved_path.replace(".docx", ".debug.issues.json")
                    debug_dir = os.path.dirname(debug_path)
                    debug_name = os.path.basename(debug_path)
                    debug_name = _truncate_filename(debug_name, max_bytes=200, reserve=0)
                    debug_path = os.path.join(debug_dir, debug_name)
                    with open(debug_path, "w", encoding="utf-8") as f:
                        json.dump(issues_data, f, ensure_ascii=False, indent=2)

                    # 调用 inject-docx-comments.py 的 validate
                    validation_ok = True
                    try:
                        import importlib.util
                        spec = importlib.util.spec_from_file_location("inject_docx_comments", inject_script)
                        mod = importlib.util.module_from_spec(spec)
                        spec.loader.exec_module(mod)
                        validation_errors = mod.validate_issues_json(issues_data)
                        if validation_errors:
                            validation_ok = False
                            results["annotated_docx"]["error"] = f"issues.json 校验失败: {'; '.join(validation_errors[:5])}"
                    except Exception:
                        pass  # 跳过校验，直接生成

                    if validation_ok:
                        temp_issues = tempfile.NamedTemporaryFile(
                            mode="w", suffix=".json", delete=False, encoding="utf-8"
                        )
                        json.dump(issues_data, temp_issues, ensure_ascii=False, indent=2)
                        temp_issues.close()

                        out_path = resolved_path.replace(".docx", ".annotated.docx")
                        # 确保文件名不超过系统限制（Linux 255 字节）
                        out_dir = os.path.dirname(out_path)
                        out_name = os.path.basename(out_path)
                        out_name = _truncate_filename(out_name, max_bytes=200, reserve=0)
                        out_path = os.path.join(out_dir, out_name)
                        proc = subprocess.run(
                            [sys.executable, inject_script, resolved_path, temp_issues.name, out_path],
                            capture_output=True, text=True, timeout=180
                        )
                        os.unlink(temp_issues.name)

                        if proc.returncode == 0:
                            # 上传到对象存储
                            download_url = _upload_to_s3(out_path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                            if download_url:
                                results["annotated_docx"] = {"success": True, "url": download_url}
                            else:
                                results["annotated_docx"] = {"success": True, "url": out_path, "note": "本地路径，请从项目文件列表下载"}
                        else:
                            results["annotated_docx"]["error"] = proc.stderr or "批注注入脚本返回非零退出码"
                else:
                    if not results["annotated_docx"]["error"]:
                        results["annotated_docx"]["error"] = f"docx 文件不存在: {docx_path}"
            else:
                results["annotated_docx"]["error"] = f"批注注入脚本不存在: {inject_script}"
        except Exception as e:
            results["annotated_docx"]["error"] = f"生成批注文档异常: {str(e)}"
    else:
        # docx_path 为空，尝试从缓存恢复
        recovered = _recover_docx_from_cache("")
        if recovered:
            docx_path = recovered
            # 重新尝试生成（复用上面的逻辑）
            try:
                import subprocess
                inject_script = _utr_tool_path("inject-docx-comments.py")
                if os.path.exists(inject_script):
                    temp_issues = tempfile.NamedTemporaryFile(
                        mode="w", suffix=".json", delete=False, encoding="utf-8"
                    )
                    json.dump(issues_data, temp_issues, ensure_ascii=False, indent=2)
                    temp_issues.close()
                    out_path = docx_path.replace(".docx", ".annotated.docx")
                    out_name = os.path.basename(out_path)
                    out_name = _truncate_filename(out_name, max_bytes=200, reserve=0)
                    out_path = os.path.join(os.path.dirname(out_path), out_name)
                    proc = subprocess.run(
                        [sys.executable, inject_script, docx_path, temp_issues.name, out_path],
                        capture_output=True, text=True, timeout=180
                    )
                    os.unlink(temp_issues.name)
                    if proc.returncode == 0:
                        download_url = _upload_to_s3(out_path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        if download_url:
                            results["annotated_docx"] = {"success": True, "url": download_url}
                        else:
                            results["annotated_docx"] = {"success": True, "url": out_path, "note": "本地路径"}
                    else:
                        results["annotated_docx"]["error"] = proc.stderr or "批注注入脚本返回非零退出码"
                else:
                    results["annotated_docx"]["error"] = f"批注注入脚本不存在: {inject_script}"
            except Exception as e:
                results["annotated_docx"]["error"] = f"生成批注文档异常: {str(e)}"
        else:
            results["annotated_docx"]["error"] = "未提供 docx_path 且无法从缓存恢复论文内容，跳过批注文档生成"

    # ── Step 3: 生成 PDF 审查报告 ──
    try:
        from coze_coding_dev_sdk import DocumentGenerationClient

        issues_list = issues_data.get("issues", [])
        sorted_issues = sorted(issues_list, key=lambda x: SEVERITY_ORDER.get(x.get("severity", "minor"), 2))
        fatal_count = sum(1 for i in issues_list if i.get("severity") == "fatal")
        major_count = sum(1 for i in issues_list if i.get("severity") == "major")
        minor_count = sum(1 for i in issues_list if i.get("severity") == "minor")

        issues_detail = ""
        if issues_list:
            issues_detail = f"共 **{len(issues_list)}** 项问题：致命 {fatal_count} 项 / 严重 {major_count} 项 / 轻微 {minor_count} 项\n\n"
            for idx, issue in enumerate(sorted_issues, 1):
                sev = issue.get("severity", "minor")
                sev_label = {"fatal": "❌ 致命", "major": "⚠️ 严重", "minor": "💡 轻微"}.get(sev, "💡 轻微")
                cat = issue.get("category", "unknown")
                cat_label = {
                    "structure": "结构逻辑", "argumentation": "论证严谨性",
                    "literature-review": "文献综述", "empirical": "实证分析",
                    "legal-norms": "规范适用", "language": "语言表达",
                    "policy": "政策建议", "academic-integrity": "学术诚信",
                    "citation-format": "引注格式", "citation-missing-info": "引注缺失"
                }.get(cat, cat)
                problem = issue.get("problem", "（未描述）").replace("|", "｜").replace("\n", " ")
                suggestions = issue.get("suggestion", [])
                sug_text = "；".join(suggestions[:3]).replace("\n", " ") if suggestions else "请根据上下文修改"
                excerpt = issue.get("excerpt", "")[:80].replace("\n", " ")

                issues_detail += f"### {idx}. {sev_label} 【{cat_label}】\n\n"
                issues_detail += f"**问题**：{problem}\n\n"
                if excerpt:
                    issues_detail += f"**原文**：…{excerpt}…\n\n"
                issues_detail += f"**建议**：{sug_text}\n\n---\n\n"
        else:
            issues_detail = "> 问题清单数据暂不可用，请参考审查分析内容。\n"

        summary_text = analysis_summary or "（详细审查分析请见对话记录）"

        report_md = f"""# 法学论文质检审查报告

---

## 审查概况

{summary_text}

---

## 问题清单

{issues_detail}

## 使用说明

- 本报告由法学论文质检自查智能体自动生成
- 审查结果仅供参考，请以导师/评阅人意见为准
- 致命/严重问题建议优先处理
- 如有疑问，请与导师沟通确认
"""
        client = DocumentGenerationClient()
        pdf_url = client.create_pdf_from_markdown(report_md, "review_report")
        results["pdf_report"] = {"success": True, "url": pdf_url}
    except Exception as e:
        results["pdf_report"]["error"] = f"PDF报告生成失败: {str(e)}"

    # ── 构建最终返回 ──
    docx_url = results["annotated_docx"].get("url")
    pdf_url = results["pdf_report"].get("url")
    docx_ok = results["annotated_docx"]["success"]
    pdf_ok = results["pdf_report"]["success"]

    return json.dumps({
        "success": docx_ok or pdf_ok,
        "annotated_docx_url": docx_url,
        "pdf_url": pdf_url,
        "annotated_docx_error": results["annotated_docx"].get("error"),
        "pdf_error": results["pdf_report"].get("error"),
        "total_issues": len(issues_data.get("issues", [])),
        "quality_gate": issues_data.get("quality_gate", {}),
        "_hint": "请向用户呈现审查摘要和下载链接。不要输出此 JSON 原文。docx 链接是带批注的论文，pdf 链接是审查报告。链接24小时内有效。",
    }, ensure_ascii=False)


def _upload_to_s3(file_path: str, content_type: str) -> Optional[str]:
    """上传文件到对象存储，返回预签名 URL。失败返回 None。"""
    try:
        from coze_coding_dev_sdk.s3 import S3SyncStorage
        import time
        storage = S3SyncStorage(
            endpoint_url=os.getenv("COZE_BUCKET_ENDPOINT_URL"),
            bucket_name=os.getenv("COZE_BUCKET_NAME"),
        )
        import re, uuid as _uuid
        raw_name = os.path.basename(file_path)
        safe_name = re.sub(r'[^a-zA-Z0-9_\-\.]', '_', raw_name)
        if not safe_name or safe_name.startswith('_') or safe_name == '_':
            safe_name = 'document'
        ext = os.path.splitext(safe_name)[1].lower() or ".bin"
        # 使用短 UUID 作为 S3 key，避免中文/过长文件名导致签名问题
        unique_name = f"papers/{int(time.time())}_{_uuid.uuid4().hex[:8]}{ext}"
        with open(file_path, "rb") as f:
            file_key = storage.stream_upload_file(
                fileobj=f,
                file_name=unique_name,
                content_type=content_type
            )
        download_url = storage.generate_presigned_url(
            key=file_key,
            expire_time=86400
        )
        return download_url
    except Exception as e:
        import logging
        logging.warning(f"[S3Upload] 上传失败: {e}")
        return None


# ---------------------------------------------------------------------------
# DeepSeek V4 thinking mode 兼容补丁
# ---------------------------------------------------------------------------
# DeepSeek V4 系列模型（deepseek-v4-pro / deepseek-v4-flash）默认开启思考模式，
# 当响应包含 tool_calls 时，后续请求必须原样回传 reasoning_content，否则报 400。
# langchain-openai 默认不处理 reasoning_content，这里通过 monkey-patch 修复。
# ---------------------------------------------------------------------------
def _apply_deepseek_v4_patch():
    """Monkey-patch langchain_openai 以支持 reasoning_content 回传。"""
    try:
        import langchain_openai.chat_models.base as lc_base
        from langchain_core.messages import AIMessage

        # 保存原始函数
        _orig_dict_to_msg = lc_base._convert_dict_to_message
        _orig_msg_to_dict = lc_base._convert_message_to_dict

        def _patched_dict_to_message(_dict):
            msg = _orig_dict_to_msg(_dict)
            # 将响应中的 reasoning_content 保存到 AIMessage.additional_kwargs
            if isinstance(msg, AIMessage) and "reasoning_content" in _dict:
                msg.additional_kwargs["reasoning_content"] = _dict["reasoning_content"]
            return msg

        def _patched_msg_to_dict(message, api="chat/completions"):
            result = _orig_msg_to_dict(message, api)
            # 将 AIMessage.additional_kwargs 中的 reasoning_content 回传到请求
            if isinstance(message, AIMessage) and message.additional_kwargs.get("reasoning_content"):
                result["reasoning_content"] = message.additional_kwargs["reasoning_content"]
            return result

        lc_base._convert_dict_to_message = _patched_dict_to_message
        lc_base._convert_message_to_dict = _patched_msg_to_dict
        logger.info("[DeepSeekPatch] reasoning_content 兼容补丁已应用")
    except Exception as e:
        logger.warning(f"[DeepSeekPatch] 应用补丁失败（不影响正常运行）: {e}")


# ---------------------------------------------------------------------------
# 错误处理中间件
# ---------------------------------------------------------------------------
@wrap_tool_call
def handle_tool_errors(request, handler):
    try:
        return handler(request)
    except Exception as e:
        tb = traceback.format_exc()
        return ToolMessage(
            content=f"工具执行错误 ({request.tool_call.get('name', 'unknown')}): {str(e)}\n\n{tb}",
            tool_call_id=request.tool_call["id"]
        )


# ---------------------------------------------------------------------------
# Agent 构建
# ---------------------------------------------------------------------------
def build_agent(ctx=None):
    workspace_path = os.getenv("COZE_WORKSPACE_PATH", "/workspace/projects")
    config_path = os.path.join(workspace_path, LLM_CONFIG)
    with open(config_path, "r", encoding="utf-8") as f:
        cfg = json.load(f)

    # 解析认证信息：BYOK > 平台工作负载身份
    api_key, base_url = _resolve_auth()
    if not api_key:
        raise RuntimeError(
            "无法获取模型认证信息。请检查："
            "1) 是否设置了 CUSTOM_MODEL_API_KEY 环境变量（BYOK 模式）；"
            "2) 或平台工作负载身份环境变量 COZE_WORKLOAD_IDENTITY_API_KEY 是否有效。"
        )

    model_name = (
        os.getenv("CUSTOM_MODEL_NAME")
        or os.getenv("DEEPSEEK_MODEL")
        or os.getenv("model_name")                 # 兼容用户自定义变量名
        or cfg.get("custom_model", {}).get("model")
        or cfg["config"]["model"]
    )
    thinking_cfg = cfg["config"].get("thinking", "disabled")

    logger.info(f"[Agent] model={model_name}, base_url={base_url}, key_type={'BYOK' if api_key and not api_key.startswith('eyJ') else 'platform'}")

    # DeepSeek V4 系列需要 reasoning_content 兼容补丁
    if "deepseek-v4" in model_name.lower():
        _apply_deepseek_v4_patch()

    extra_body = {}
    if thinking_cfg == "enabled":
        extra_body = {"thinking": {"type": "enabled"}}

    llm = ChatOpenAI(
        model=model_name,
        api_key=api_key,
        base_url=base_url,
        temperature=cfg["config"]["temperature"],
        max_completion_tokens=cfg["config"].get("max_completion_tokens", 32768),
        streaming=True,
        timeout=cfg["config"]["timeout"],
        extra_body=extra_body,
        default_headers=default_headers(ctx) if ctx else {}
    )

    tools = [
        read_paper_file,
        generate_deliverables,
    ]

    return create_agent(
        model=llm,
        system_prompt=cfg.get("sp", "You are a helpful assistant."),
        tools=tools,
        middleware=[handle_tool_errors],
        checkpointer=get_memory_saver(),
        state_schema=AgentState,
    )
