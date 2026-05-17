#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档文本抽取脚本 (extract-docx.py)

从Word文档(.docx)中提取文本内容，保留结构信息。
支持提取脚注、尾注、批注等。
"""

import argparse
import json
import logging
import re
import sys
from pathlib import Path
from typing import Dict, List, Any, Optional

try:
    from docx import Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph
except ImportError:
    print("错误: 需要安装 python-docx 库")
    print("运行: uv add python-docx")
    sys.exit(1)


logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class DOCXTextExtractor:
    """Word文档文本提取器"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = None  # type: ignore
        self.metadata: Dict[str, Any] = {}
        
    def open(self) -> bool:
        """打开Word文档"""
        try:
            self.doc = Document(self.file_path)
            self._extract_metadata()
            return True
        except Exception as e:
            logger.error(f"无法打开Word文档: {e}")
            return False
    
    def _extract_metadata(self):
        """提取文档元数据"""
        if self.doc:
            core_props = self.doc.core_properties
            self.metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or "",
            }
    
    def extract_text(self) -> str:
        """
        提取文档全文文本
        
        Returns:
            text: 文档文本内容
        """
        if not self.doc:
            return ""
        
        paragraphs = []
        
        # 提取正文段落
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if text:
                # 添加段落样式信息
                style_name = para.style.name if para.style else "Normal"
                paragraphs.append(f"[{style_name}]\n{text}")
        
        # 提取表格内容
        for table in self.doc.tables:
            table_text = self._extract_table_text(table)
            if table_text:
                paragraphs.append(f"[表格]\n{table_text}")
        
        return "\n\n".join(paragraphs)
    
    def _extract_table_text(self, table: Table) -> str:
        """提取表格文本"""
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(" | ".join(cells))
        return "\n".join(rows_data)
    
    def extract_with_positions(self) -> List[Dict[str, Any]]:
        """
        提取文本并记录位置信息（用于批注定位）
        
        Returns:
            List[Dict]: 每段的文本和位置信息
        """
        if not self.doc:
            return []
        
        blocks = []
        current_pos = 0
        
        # 处理段落
        for para in self.doc.paragraphs:
            text = para.text
            if text.strip():
                blocks.append({
                    "type": "paragraph",
                    "text": text,
                    "style": para.style.name if para.style else "Normal",
                    "position": current_pos,
                    "length": len(text),
                    "runs": [
                        {
                            "text": run.text,
                            "bold": run.bold,
                            "italic": run.italic,
                        }
                        for run in para.runs if run.text.strip()
                    ]
                })
                current_pos += len(text) + 2
        
        # 处理表格
        for table in self.doc.tables:
            table_text = self._extract_table_text(table)
            if table_text.strip():
                blocks.append({
                    "type": "table",
                    "text": table_text,
                    "position": current_pos,
                    "length": len(table_text),
                })
                current_pos += len(table_text) + 2
        
        return blocks
    
    def extract_structured(self) -> Dict[str, Any]:
        """
        提取结构化信息
        
        Returns:
            Dict: 包含文本、结构和元数据的完整报告
        """
        if not self.doc:
            return {"error": "文档未打开"}
        
        text = self.extract_text()
        blocks = self.extract_with_positions()
        
        result = {
            "metadata": self.metadata,
            "text": text,
            "blocks": blocks,
            "structure": self._analyze_structure(text),
            "footnotes": self._extract_footnotes(),
            "endnotes": self._extract_endnotes(),
            "comments": self._extract_comments(),
            "references": self._extract_references(text),
            "legal_citations": self._extract_legal_citations(text),
            "case_citations": self._extract_case_citations(text),
        }
        
        return result
    
    def _analyze_structure(self, text: str) -> Dict[str, Any]:
        """分析文档结构"""
        structure = {
            "title": "",
            "chapters": [],
            "sections": [],
            "total_headings": 0,
            "word_count": len(text),
            "char_count": len(text.replace(" ", "")),
        }
        
        lines = text.split('\n')
        
        # 匹配标题样式
        heading_styles = [
            "Heading 1", "Heading 2", "Heading 3",
            "标题 1", "标题 2", "标题 3",
            "标题", "子标题"
        ]
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # 检查是否为标题（通过样式标记）
            if line.startswith("[") and "]" in line:
                style = line[1:line.index("]")]
                if style in heading_styles:
                    structure["total_headings"] += 1
                    if not structure["title"]:
                        structure["title"] = line[line.index("]")+1:].strip()
                    elif "Heading 1" in style or "标题 1" in style:
                        structure["chapters"].append({
                            "title": line[line.index("]")+1:].strip(),
                            "line_num": i
                        })
                    else:
                        structure["sections"].append({
                            "title": line[line.index("]")+1:].strip(),
                            "line_num": i
                        })
            else:
                # 通过内容模式检测标题
                heading_patterns = [
                    r'^第[一二三四五六七八九十百\d]+[章节部篇]',
                    r'^[一二三四五六七八九十百\d]+[、\.](?!\d)',
                    r'^(?:摘要|Abstract|引言|Introduction|结论|Conclusion|参考文献|致谢)',
                ]
                for pattern in heading_patterns:
                    if re.match(pattern, line):
                        structure["total_headings"] += 1
                        if not structure["title"]:
                            structure["title"] = line[:100]
                        else:
                            structure["sections"].append({
                                "title": line,
                                "line_num": i
                            })
                        break
        
        return structure
    
    def _extract_footnotes(self) -> List[Dict[str, Any]]:
        """提取脚注"""
        footnotes = []
        
        # 尝试从XML中提取脚注
        try:
            part = self.doc.part
            footnotes_part = getattr(part, 'footnotes_part', None)
            if footnotes_part is not None:
                footnotes_xml = footnotes_part._element
                for fn in footnotes_xml.findall('.//{*}footnote'):
                    fn_id = fn.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    fn_text = ''.join([t.text for t in fn.iter() if t.text])
                    if fn_text.strip():
                        footnotes.append({
                            "id": fn_id,
                            "text": fn_text[:200]
                        })
        except Exception as e:
            logger.debug(f"无法提取脚注: {e}")
        
        # 如果XML提取失败，从正文中提取
        if not footnotes:
            text = self.extract_text()
            footnote_pattern = r'^\[\d+\]|\^\d+'
            for match in re.finditer(footnote_pattern, text):
                start = match.start()
                end = min(start + 200, len(text))
                context = text[start:end]
                footnotes.append({
                    "text": context,
                    "position": start
                })
        
        return footnotes[:100]
    
    def _extract_endnotes(self) -> List[Dict[str, str]]:
        """提取尾注"""
        endnotes = []
        
        try:
            part = self.doc.part
            endnotes_part = getattr(part, 'endnotes_part', None)
            if endnotes_part is not None:
                endnotes_xml = endnotes_part._element
                for en in endnotes_xml.findall('.//{*}endnote'):
                    en_id = en.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    en_text = ''.join([t.text for t in en.iter() if t.text])
                    if en_text.strip():
                        endnotes.append({
                            "id": en_id,
                            "text": en_text[:200]
                        })
        except Exception as e:
            logger.debug(f"无法提取尾注: {e}")
        
        return endnotes
    
    def _extract_comments(self) -> List[Dict[str, Any]]:
        """提取批注"""
        comments = []
        
        try:
            part = self.doc.part
            comments_part = getattr(part, 'comments_part', None)
            if comments_part is not None:
                comments_xml = comments_part._element
                for comment in comments_xml.findall('.//{*}comment'):
                    comment_id = comment.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    author = comment.get('author', 'Unknown')
                    date = comment.get('date', '')
                    text = ''.join([t.text for t in comment.iter() if t.text])
                    if text.strip():
                        comments.append({
                            "id": comment_id,
                            "author": author,
                            "date": date,
                            "text": text
                        })
        except Exception as e:
            logger.debug(f"无法提取批注: {e}")
        
        return comments
    
    def _extract_references(self, text: str) -> List[Dict[str, str]]:
        """提取参考文献"""
        references = []
        
        # 查找参考文献章节
        ref_section_match = re.search(
            r'(?:参考文献|References|引用文献)[：:]\s*\n(.*?)(?=\n(?:脚注|注释|正文|$))',
            text,
            re.DOTALL | re.IGNORECASE
        )
        
        if ref_section_match:
            ref_text = ref_section_match.group(1)
            lines = ref_text.split('\n')
            
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue
                
                # 参考文献通常以序号开头
                ref_match = re.match(r'^\[\d+\]|\d+\.|^[\[（\(]\d+[\]）\)]', line)
                if ref_match:
                    references.append({
                        "text": line[:300],
                        "type": self._detect_ref_type(line)
                    })
        
        return references[:100]
    
    def _detect_ref_type(self, text: str) -> str:
        """检测参考文献类型"""
        if re.search(r'[《》\[\(（]法|\[法\]|F\d+', text):
            return "law"
        if re.search(r'\d{4}[年]?\s*[第卷期]?\s*\d+|\(\d{4}\)', text):
            return "journal"
        if re.search(r'(?:北京|上海|广州|人民出版社|法律出版社)', text):
            return "book"
        if re.search(r'(?:博士|硕士)学位论文', text):
            return "thesis"
        if re.search(r'http[s]?://|www\.', text):
            return "web"
        if re.search(r'[A-Z][a-z]+,\s*[A-Z]\.', text):
            return "foreign"
        return "other"
    
    def _extract_legal_citations(self, text: str) -> List[Dict[str, str]]:
        """提取法律条文引用"""
        citations = []
        
        legal_patterns = [
            r'《([^《》]+)》第(\d+)[条款节章部]',
            r'(?:第)(\d+)(?:条|款|节)(?:之(\d+))?',
            r'《([^《》]+)》第(\d+)条(?:\s*第(\d+)款)?',
        ]
        
        for pattern in legal_patterns:
            for match in re.finditer(pattern, text):
                citations.append({
                    "text": match.group(0),
                    "law": match.group(1) if match.lastindex >= 1 else "",
                    "article": match.group(2) if match.lastindex >= 2 else "",
                    "paragraph": match.group(3) if match.lastindex >= 3 else "",
                })
        
        return citations[:50]
    
    def _extract_case_citations(self, text: str) -> List[Dict[str, str]]:
        """提取案例引用"""
        citations = []
        
        case_patterns = [
            r'([^\s]+)\s*案\s*(?:第?\s*(\d+)\s*号)?',
            r'(?:最高法|最高人民法院|最高检|最高人民检察院)(?:民|刑|行|知)(?:终|再|二)?(?:(?:字)|)\s*第?\s*(\d+)\s*号',
            r'\[(\d{4})\]\s*(?:最高法|最高人民法院)(?:民|刑|行|知).*?第?\s*(\d+)\s*号',
        ]
        
        for pattern in case_patterns:
            for match in re.finditer(pattern, text):
                citations.append({
                    "text": match.group(0),
                    "year": match.group(1) if match.lastindex >= 1 else "",
                    "case_num": match.group(2) if match.lastindex >= 2 else "",
                })
        
        return citations[:30]
    
    def get_paragraph_count(self) -> int:
        """获取段落数量"""
        if self.doc:
            return len(self.doc.paragraphs)
        return 0
    
    def get_table_count(self) -> int:
        """获取表格数量"""
        if self.doc:
            return len(self.doc.tables)
        return 0
    
    def close(self):
        """关闭文档"""
        self.doc = None
    
    def __enter__(self):
        self.open()
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()
        return False


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description="从Word文档中提取文本内容，支持结构化分析"
    )
    parser.add_argument(
        "input_file",
        help="输入Word文档路径"
    )
    parser.add_argument(
        "-o", "--output",
        help="输出文本文件路径（可选）"
    )
    parser.add_argument(
        "-j", "--json",
        action="store_true",
        help="输出JSON格式的完整报告"
    )
    parser.add_argument(
        "-s", "--structured",
        action="store_true",
        help="输出结构化分析"
    )
    parser.add_argument(
        "-p", "--positions",
        action="store_true",
        help="输出带位置信息的结构化数据"
    )
    parser.add_argument(
        "-q", "--quiet",
        action="store_true",
        help="静默模式"
    )
    
    args = parser.parse_args()
    
    if not Path(args.input_file).exists():
        print(f"错误: 文件不存在: {args.input_file}", file=sys.stderr)
        sys.exit(1)
    
    with DOCXTextExtractor(args.input_file) as extractor:
        if args.structured or args.json:
            result = extractor.extract_structured()
        elif args.positions:
            result = {
                "blocks": extractor.extract_with_positions(),
                "metadata": extractor.metadata
            }
        else:
            text = extractor.extract_text()
            result = {
                "text": text,
                "metadata": extractor.metadata,
                "paragraph_count": extractor.get_paragraph_count(),
                "table_count": extractor.get_table_count()
            }
    
    # 输出结果
    if args.json or args.structured:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print(result["text"])
        
        if not args.quiet:
            print("\n" + "="*60)
            print("文档信息")
            print("="*60)
            meta = result.get("metadata", {})
            print(f"标题: {meta.get('title', 'N/A')}")
            print(f"作者: {meta.get('author', 'N/A')}")
            print(f"段落数: {result.get('paragraph_count', 'N/A')}")
            print(f"表格数: {result.get('table_count', 'N/A')}")
    
    # 保存到文件
    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            if args.json or args.structured:
                json.dump(result, f, ensure_ascii=False, indent=2)
            else:
                f.write(result["text"])
        
        if not args.quiet:
            print(f"\n✓ 文本已保存到: {args.output}")


if __name__ == "__main__":
    main()
