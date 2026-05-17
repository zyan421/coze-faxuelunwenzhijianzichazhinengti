#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word批注注入脚本 (inject-docx-comments.py)

向Word文档中注入批注，支持：
1. 按文本位置定位并添加批注
2. 按段落/行号添加批注
3. 高亮标注问题文字
4. 支持多种批注格式
"""

import argparse
import json
import logging
import re
import sys
from dataclasses import dataclass, asdict
from enum import Enum
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple

try:
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("错误: 需要安装 python-docx 库")
    print("运行: uv add python-docx")
    sys.exit(1)


logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class Severity(Enum):
    """问题严重程度"""
    FATAL = "fatal"        # 致命问题
    MAJOR = "major"        # 重大问题
    MINOR = "minor"        # 一般问题
    INFO = "info"          # 提示信息


@dataclass
class Annotation:
    """批注对象"""
    id: str
    severity: Severity
    issue_type: str
    description: str
    location: str          # 位置描述
    suggestion: str        # 修改建议
    highlight_text: str = ""  # 需要高亮的文本
    evidence: str = ""     # 证据
    
    def to_dict(self) -> Dict[str, Any]:
        result = asdict(self)
        result['severity'] = self.severity.value
        return result


@dataclass
class AnnotationResult:
    """批注结果"""
    success: bool
    output_path: str
    annotations_added: int
    annotations_failed: int
    failed_annotations: List[Dict] = None
    error_message: str = ""
    
    def to_dict(self) -> Dict[str, Any]:
        result = asdict(self)
        if self.failed_annotations is None:
            result['failed_annotations'] = []
        return result


class DOCXCommentInjector:
    """Word批注注入器"""
    
    # 严重程度对应的颜色
    SEVERITY_COLORS = {
        Severity.FATAL: RGBColor(0xFF, 0x00, 0x00),    # 红色
        Severity.MAJOR: RGBColor(0xFF, 0x66, 0x00),    # 橙色
        Severity.MINOR: RGBColor(0xFF, 0xCC, 0x00),    # 黄色
        Severity.INFO: RGBColor(0x00, 0x00, 0xFF),     # 蓝色
    }
    
    # 严重程度前缀
    SEVERITY_PREFIX = {
        Severity.FATAL: "【致命】",
        Severity.MAJOR: "【重大】",
        Severity.MINOR: "【建议】",
        Severity.INFO: "【提示】",
    }
    
    def __init__(self, input_path: str, output_path: str):
        self.input_path = input_path
        self.output_path = output_path
        self.doc = None  # type: ignore
        self.annotations: List[Annotation] = []
        
    def load(self) -> bool:
        """加载文档"""
        try:
            self.doc = Document(self.input_path)
            return True
        except Exception as e:
            logger.error(f"无法加载文档: {e}")
            return False
    
    def add_annotation(self, annotation: Annotation) -> bool:
        """
        添加单条批注
        
        Args:
            annotation: 批注对象
            
        Returns:
            是否成功添加
        """
        if not self.doc:
            logger.error("文档未加载")
            return False
        
        try:
            # 尝试多种方式定位
            added = False
            
            # 方式1: 通过高亮文本定位
            if annotation.highlight_text:
                added = self._add_annotation_by_text(annotation)
            
            # 方式2: 通过段落索引定位（如果有）
            if not added and annotation.location:
                # 尝试解析位置
                location_info = self._parse_location(annotation.location)
                if location_info:
                    added = self._add_annotation_by_location(annotation, location_info)
            
            # 方式3: 在文档末尾添加通用批注
            if not added:
                added = self._add_annotation_at_end(annotation)
            
            if added:
                self.annotations.append(annotation)
            
            return added
            
        except Exception as e:
            logger.error(f"添加批注失败: {e}")
            return False
    
    def _parse_location(self, location: str) -> Optional[Dict[str, Any]]:
        """解析位置字符串"""
        info = {}
        
        # 解析 "第X页，第Y行"
        page_match = re.search(r'第(\d+)页', location)
        if page_match:
            info['page'] = int(page_match.group(1))
        
        line_match = re.search(r'第(\d+)行', location)
        if line_match:
            info['line'] = int(line_match.group(1))
        
        # 解析 "第X段"
        para_match = re.search(r'第(\d+)段', location)
        if para_match:
            info['paragraph'] = int(para_match.group(1))
        
        # 解析 "第X章第Y节"
        chapter_match = re.search(r'第([一二三四五六七八九十百\d]+)章', location)
        if chapter_match:
            info['chapter'] = self._chinese_to_arabic(chapter_match.group(1))
        
        return info if info else None
    
    @staticmethod
    def _chinese_to_arabic(cn: str) -> int:
        """中文数字转阿拉伯数字"""
        CN_NUM = {
            '零': 0, '一': 1, '二': 2, '三': 3, '四': 4,
            '五': 5, '六': 6, '七': 7, '八': 8, '九': 9,
            '十': 10, '百': 100
        }
        
        result = 0
        temp = 0
        
        for char in cn:
            if char == '十' or char == '百':
                if temp == 0:
                    temp = 1
                result = result + temp * CN_NUM.get(char, 0)
                temp = 0
            else:
                temp = temp * 10 + CN_NUM.get(char, 0)
        
        result += temp
        return result if result > 0 else 0
    
    def _add_annotation_by_text(self, annotation: Annotation) -> bool:
        """通过文本内容定位并添加批注"""
        highlight_text = annotation.highlight_text.strip()
        if not highlight_text:
            return False
        
        # 在所有段落中查找匹配文本
        for para_idx, para in enumerate(self.doc.paragraphs):
            para_text = para.text
            if highlight_text in para_text:
                # 找到匹配
                return self._insert_comment_in_paragraph(
                    para, highlight_text, annotation
                )
        
        # 在表格中查找
        for table in self.doc.tables:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    if highlight_text in cell.text:
                        # 找到匹配 - 在表格单元格中添加批注
                        para = cell.paragraphs[0]
                        return self._insert_comment_in_paragraph(
                            para, highlight_text, annotation
                        )
        
        return False
    
    def _insert_comment_in_paragraph(
        self, para, highlight_text: str, annotation: Annotation
    ) -> bool:
        """在段落中插入批注"""
        try:
            # 获取段落XML
            p = para._p
            
            # 创建批注XML元素
            comment = self._create_comment_element(annotation)
            
            # 在匹配文本后插入批注引用
            runs = para.runs
            if runs:
                # 在第一个匹配的run后插入批注
                first_run = runs[0]
                # 添加高亮
                self._add_highlight(first_run, annotation.severity)
                
                # 在run后插入批注
                r = first_run._r
                r.addnext(comment)
            else:
                # 如果没有runs，在段落末尾添加
                p.append(comment)
            
            return True
            
        except Exception as e:
            logger.debug(f"插入批注失败: {e}")
            return False
    
    def _create_comment_element(self, annotation: Annotation):
        """创建批注XML元素"""
        # 创建commentRangeStart
        start = OxmlElement('w:commentRangeStart')
        start.set(qn('w:id'), str(hash(annotation.id) % 100000))
        
        # 创建commentRangeEnd
        end = OxmlElement('w:commentRangeEnd')
        end.set(qn('w:id'), str(hash(annotation.id) % 100000))
        
        # 创建commentReference
        ref = OxmlElement('w:commentReference')
        ref.set(qn('w:id'), str(hash(annotation.id) % 100000))
        
        # 创建run包含引用
        r = OxmlElement('w:r')
        r.append(ref)
        
        return r
    
    def _add_highlight(self, run, severity: Severity) -> bool:
        """添加高亮"""
        try:
            # 设置背景色
            rPr = run._r.get_or_add_rPr()
            highlight = OxmlElement('w:highlight')
            highlight.set(qn('w:val'), self._severity_to_highlight_color(severity))
            rPr.append(highlight)
            return True
        except Exception as e:
            logger.debug(f"添加高亮失败: {e}")
            return False
    
    @staticmethod
    def _severity_to_highlight_color(severity: Severity) -> str:
        """严重程度对应的高亮颜色"""
        colors = {
            Severity.FATAL: "red",
            Severity.MAJOR: "orange",
            Severity.MINOR: "yellow",
            Severity.INFO: "cyan",
        }
        return colors.get(severity, "yellow")
    
    def _add_annotation_by_location(
        self, annotation: Annotation, location: Dict[str, Any]
    ) -> bool:
        """通过位置信息添加批注"""
        target_para_idx = None
        
        # 通过段落索引
        if 'paragraph' in location:
            target_para_idx = location['paragraph'] - 1  # 转0-based
        elif 'page' in location and 'line' in location:
            # 通过页码和行号定位（粗略估计每页约40行）
            page = location['page']
            line = location.get('line', 1)
            target_para_idx = (page - 1) * 40 + line - 1
        
        if target_para_idx is not None and target_para_idx < len(self.doc.paragraphs):
            para = self.doc.paragraphs[target_para_idx]
            if para.text:
                return self._insert_comment_in_paragraph(
                    para, para.text[:50], annotation
                )
        
        return False
    
    def _add_annotation_at_end(self, annotation: Annotation) -> bool:
        """在文档末尾添加批注"""
        try:
            # 添加一个批注段落
            prefix = self.SEVERITY_PREFIX.get(annotation.severity, "")
            comment_text = f"\n{prefix}{annotation.issue_type}: {annotation.description}"
            
            if annotation.suggestion:
                comment_text += f"\n建议: {annotation.suggestion}"
            
            # 添加到文档末尾
            self.doc.add_paragraph(comment_text)
            return True
            
        except Exception as e:
            logger.error(f"在末尾添加批注失败: {e}")
            return False
    
    def add_comments_from_issues(
        self, issues: List[Dict]
    ) -> Tuple[int, int, List[Dict]]:
        """
        从issues列表批量添加批注
        
        Args:
            issues: 问题列表，每项包含severity, type, description等
            
        Returns:
            (成功数, 失败数, 失败列表)
        """
        success_count = 0
        failed_count = 0
        failed_list = []
        
        for idx, issue in enumerate(issues):
            try:
                severity = Severity(issue.get('severity', 'minor'))
            except ValueError:
                severity = Severity.INFO
            
            annotation = Annotation(
                id=f"issue_{idx}",
                severity=severity,
                issue_type=issue.get('type', issue.get('issue_type', 'general')),
                description=issue.get('description', ''),
                location=issue.get('location', ''),
                suggestion=issue.get('suggestion', ''),
                highlight_text=issue.get('highlight_text', ''),
                evidence=issue.get('evidence', ''),
            )
            
            if self.add_annotation(annotation):
                success_count += 1
            else:
                failed_count += 1
                failed_list.append({
                    'index': idx,
                    'issue': issue,
                    'reason': '无法定位到文档位置'
                })
        
        return success_count, failed_count, failed_list
    
    def save(self) -> bool:
        """保存文档"""
        if not self.doc:
            return False
        
        try:
            self.doc.save(self.output_path)
            logger.info(f"文档已保存: {self.output_path}")
            return True
        except Exception as e:
            logger.error(f"保存文档失败: {e}")
            return False
    
    def get_result(self) -> AnnotationResult:
        """获取批注结果"""
        return AnnotationResult(
            success=len(self.annotations) > 0,
            output_path=self.output_path,
            annotations_added=len(self.annotations),
            annotations_failed=0,
            failed_annotations=[],
        )


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description="向Word文档中注入批注"
    )
    parser.add_argument(
        "input_file",
        help="输入Word文档路径"
    )
    parser.add_argument(
        "-o", "--output",
        required=True,
        help="输出Word文档路径"
    )
    parser.add_argument(
        "-i", "--issues",
        required=True,
        help="问题列表JSON文件路径"
    )
    parser.add_argument(
        "-q", "--quiet",
        action="store_true",
        help="静默模式"
    )
    
    args = parser.parse_args()
    
    if not Path(args.input_file).exists():
        print(f"错误: 输入文件不存在: {args.input_file}", file=sys.stderr)
        sys.exit(1)
    
    if not Path(args.issues).exists():
        print(f"错误: 问题文件不存在: {args.issues}", file=sys.stderr)
        sys.exit(1)
    
    # 加载issues
    with open(args.issues, 'r', encoding='utf-8') as f:
        issues = json.load(f)
    
    if not issues:
        print("警告: 问题列表为空")
        sys.exit(0)
    
    # 创建注入器
    injector = DOCXCommentInjector(args.input_file, args.output)
    
    if not injector.load():
        print("错误: 无法加载文档", file=sys.stderr)
        sys.exit(1)
    
    # 批量添加批注
    success, failed, failed_list = injector.add_comments_from_issues(issues)
    
    # 保存
    if injector.save():
        result = injector.get_result()
        result.annotations_failed = failed
        result.failed_annotations = failed_list
        
        if not args.quiet:
            print(f"✓ 批注添加完成")
            print(f"  成功: {success}")
            print(f"  失败: {failed}")
            print(f"  输出: {args.output}")
        
        # 保存结果JSON
        result_path = args.output.replace('.docx', '_annotation_result.json')
        with open(result_path, 'w', encoding='utf-8') as f:
            json.dump(result.to_dict(), f, ensure_ascii=False, indent=2)
        
        if not args.quiet:
            print(f"  结果: {result_path}")
        
        sys.exit(0 if success > 0 else 1)
    else:
        print("错误: 保存文档失败", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
